#!/usr/bin/env python
# coding: utf-8

# In[1]:


### import dependencies
import pandas as pd
import numpy as np
from docx import Document
import docx
import os
import re


# In[2]:


## Gettext Function
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# In[3]:


##Gettextlist function
def getTextlist(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


# In[11]:


### Training Data files
statdir_T = '/home/user/workspace/data/training/'
filelist_T = []
for filename in os.listdir('/home/user/workspace/data/training/'):
    filelist_T.append(statdir_T+filename) 
filenames_T = pd.DataFrame(filelist_T)
filename_T=os.listdir('/home/user/workspace/data/training/')
filenames_T["FileName"] = filename_T


# In[12]:


### Validation Data files
statdir_V = '/home/user/workspace/data/eval/'
filelist_V = []
for filename in os.listdir('/home/user/workspace/data/eval/'):
    filelist_V.append(statdir_V+filename) 
filenames_V = pd.DataFrame(filelist_V)
filename_V=os.listdir('/home/user/workspace/data/eval/')
filenames_V["FileName"] = filename_V


# In[15]:


### Training file name extend 
filenames_T["Name"] = filenames_T.FileName.str.split(".",expand=True)[0]
### Validation file name extend
filenames_V["Name"] = filenames_V.FileName.str.split(".",expand=True)[0]


# In[16]:


# SQL library import 
from pandasql import sqldf
pysqldf = lambda q: sqldf(q, globals())


# In[17]:


# trainging and Validation csv files
df_train = pd.read_csv("/home/user/workspace/data/TrainingTestSet.csv")
df_val = pd.read_csv("/home/user/workspace/data/ValidationSet.csv")


# In[18]:


#rename Column names train data set
df_train = df_train.rename(columns={'File Name': 'file_name', 'Aggrement Value': 'aggrement_val',
                                    'Aggrement Start Date': 'aggrement_start_date','Aggrement End Date': 'aggrement_end_date',
                                    'Renewal Notice (Days)': 'renemal_notice_days','Party One':'party1', 'Party Two':'party2' })
## Rename column of validation dataset
df_val = df_val.rename(columns={'File Name': 'file_name', 'Aggrement Value': 'aggrement_val',
                                    'Aggrement Start Date': 'aggrement_start_date','Aggrement End Date': 'aggrement_end_date',
                                    'Renewal Notice (Days)': 'renemal_notice_days','Party One':'party1', 'Party Two':'party2' })


# In[22]:


q = """SELECT
        filenames_T.*, df_train.*
     FROM
        filenames_T, df_train
     where df_train.file_name = filenames_T.Name;"""


# In[23]:


new_train = pysqldf(q)


# In[24]:


#rename first column name in the trainign dataset
new_train = new_train.rename(columns={'0':"path" })


# In[26]:


#Training dataset preprocessing
import re
new_train.shape[0]
mytext = []
for index, item in new_train["path"].items():
    itemtext = getText(item).replace(","," ")
    itemtext2 = itemtext.replace("\n", " ")
    itemtext3 = itemtext2.replace("."," ")
    itemtext4 = re.sub(' +', ' ', itemtext3)
    mytext.append(itemtext4)


# In[27]:


# separating X and Y from training Dataset
new_train["text"] =mytext
final_df = new_train.drop(['path', "FileName", "Name", "file_name"], axis = 1) 
x_train = pd.DataFrame(final_df["text"])
y_train = final_df.drop(["text"], axis =1)


# In[28]:


from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk import pos_tag, word_tokenize
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import TfidfVectorizer


# In[29]:


# lemmatization function
def lemmatize_all(sentence):
    wnl = WordNetLemmatizer()
    for word, tag in pos_tag(word_tokenize(sentence)):
        if tag.startswith("NN"):
            yield wnl.lemmatize(word, pos='n')
        elif tag.startswith('VB'):
            yield wnl.lemmatize(word, pos='v')
        elif tag.startswith('JJ'):
            yield wnl.lemmatize(word, pos='a')
        elif tag.startswith('R'):
            yield wnl.lemmatize(word, pos='r') 
        else:
            yield word


# In[30]:


X_train1=[]
X_test1=[]
print ("Train data lemmatization begins")
for i in range(0,len(x_train)):
    X_train1.append(" ".join(lemmatize_all(str(x_train['text'][i]))))
print ("Train data lemmatization ends")


# In[38]:


## Validation dataset processing 
# rename column name
filenames_V = filenames_V.rename(columns={0:"path" })


# In[40]:


# getting the Validation document text in a field
filenames_V.shape[0]
mytext_V = []
for index, item in filenames_V["path"].items():
    itemtext = getText(item).replace(","," ")
    itemtext2 = itemtext.replace("\n", " ")
    itemtext3 = itemtext2.replace("."," ")
    itemtext4 = re.sub(' +', ' ', itemtext3)
    mytext_V.append(itemtext4)


# In[46]:


print ("Test data lemmatization begins")
for i in range (0, len(filenames_V)):
    X_test1.append(" ".join(lemmatize_all(str(filenames_V['text'][i]))))
print ("Test data lemmatization ends")


# In[48]:


from keras.models import Model
from keras.layers import Input, Dense, Embedding, SpatialDropout1D, concatenate
from keras.layers import GRU, Bidirectional, GlobalAveragePooling1D, GlobalMaxPooling1D
from keras.preprocessing import text, sequence
from keras.callbacks import Callback


# In[49]:





# In[89]:


##TFIDF matrix for training set
vec = TfidfVectorizer()
train_tfidf = vec.fit_transform(X_train1)
print("Train TF-IDF Matrix Shape: ",train_tfidf.shape)
TFIDF_Matrix = pd.DataFrame(train_tfidf.toarray(), columns=vec.get_feature_names())


# In[61]:


### Train model for Aggrement Value
from numpy import loadtxt
from keras.models import Sequential
from keras.layers import Dense
modelaggre = Sequential()
modelaggre.add(Dense(20, input_dim=TFIDF_Matrix.shape[1], activation='relu'))
modelaggre.add(Dense(15, activation='relu'))
modelaggre.add(Dense(1, activation='sigmoid'))
modelaggre.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
X = TFIDF_Matrix
aggre_val = y_train["aggrement_val"].fillna(0)
aggre_val = aggre_val/max(aggre_val)
y = aggre_val
# fit the keras model on the dataset
modelaggre.fit(X, y, epochs=150, batch_size=10)


# In[79]:


### Train model for start date
modelstartdate = Sequential()
modelstartdate.add(Dense(20, input_dim=TFIDF_Matrix.shape[1], activation='relu'))
modelstartdate.add(Dense(15, activation='relu'))
modelstartdate.add(Dense(1, activation='sigmoid'))
modelstartdate.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
X = TFIDF_Matrix
y = []
for i in y_train["aggrement_start_date"].fillna(0):
    if(i==0):
        y.append(i)
    else:
        flag = i.replace(".", "").strip()
        y.append(int(flag))
# fit the keras model on the dataset
modelstartdate.fit(X, y, epochs=150, batch_size=10)


# In[80]:


### Train model for end date
modelenddate = Sequential()
modelenddate.add(Dense(20, input_dim=TFIDF_Matrix.shape[1], activation='relu'))
modelenddate.add(Dense(15, activation='relu'))
modelenddate.add(Dense(1, activation='sigmoid'))
modelenddate.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
X = TFIDF_Matrix
y = []
for i in y_train["aggrement_end_date"].fillna(0):
    if(i==0):
        y.append(i)
    else:
        flag = i.replace(".", "").strip()
        y.append(int(flag))
# fit the keras model on the dataset
modelenddate.fit(X, y, epochs=150, batch_size=10)    


# In[82]:


### Train model for renewal days
modelrenew = Sequential()
modelrenew.add(Dense(20, input_dim=TFIDF_Matrix.shape[1], activation='relu'))
modelrenew.add(Dense(15, activation='relu'))
modelrenew.add(Dense(1, activation='sigmoid'))
modelrenew.compile(loss='binary_crossentropy', optimizer='adam', metrics=['accuracy'])
X = TFIDF_Matrix
y = []
for i in y_train["renemal_notice_days"].fillna(0):
        y.append(i)
# fit the keras model on the dataset
modelrenew.fit(X, y, epochs=150, batch_size=10)  


# In[85]:


#TFIDF for validation set 
val_tfidf = vec.fit_transform(X_test1)
print("Validation TF-IDF Matrix Shape: ",val_tfidf.shape)
TFIDF_Matrix_Val = pd.DataFrame(val_tfidf.toarray(), columns=vec.get_feature_names())


# In[107]:


new=TFIDF_Matrix.append(TFIDF_Matrix_Val)
new1 = new.iloc[:,:TFIDF_Matrix.shape[1]]
new2 = new1.iloc[TFIDF_Matrix.shape[0]:,:]


# In[108]:


y_val = new2.fillna(0)


# In[109]:


y_val.shape


# In[118]:


###Predict Aggrement Value
print(modelaggre.predict(y_val)*max(y_train["aggrement_val"].fillna(0)))
###Aggrement End date
print(modelenddate.predict(y_val))
### Predict aggrement start date
print(modelstartdate.predict(y_val))
### Predict #of  renewal days 
print(modelrenew.predict(y_val))


# In[113]:


max(aggre_val)


# In[ ]:




